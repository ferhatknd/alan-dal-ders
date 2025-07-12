#!/usr/bin/env python3
import docx
import sys

def debug_docx_structure(docx_path):
    """
    DOCX dosyasının yapısını analiz et
    """
    print(f"=== Debugging {docx_path} ===")
    
    doc = docx.Document(docx_path)
    
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Paragraph içeriğini kontrol et
    print("\n--- Paragraphs ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 5:
            print(f"Para {i}: {text[:100]}{'...' if len(text) > 100 else ''}")
    
    # Tablo içeriğini kontrol et
    print("\n--- Tables ---")
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx + 1}:")
        print(f"Dimensions: {len(table.rows)} rows x {len(table.columns)} cols")
        
        for row_idx, row in enumerate(table.rows):
            row_cells = [cell.text.strip() for cell in row.cells]
            print(f"Row {row_idx}: {row_cells}")
            
            # Öğrenme birimi, kazanım sayısı anahtar kelimelerini ara
            row_text = ' '.join(row_cells).upper()
            if any(keyword in row_text for keyword in ['KAZANIM', 'ÖĞRENME', 'BİRİM', 'SÜRE', 'SAAT']):
                print(f"  *** Found relevant keywords in row {row_idx} ***")

if __name__ == "__main__":
    docx_path = "SOSYAL MEDYA HESAP İŞLEMLERİ DBF.docx"
    debug_docx_structure(docx_path)