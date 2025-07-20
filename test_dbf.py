import os
import sys
import random
import json
from modules.oku_dbf import oku_dbf, extract_genel_kazanimlar, extract_ortam_donanimi, extract_olcme_degerlendirme, DocumentProcessor

def get_subdirectories(base_path):
    """Belirtilen yol altındaki tüm alt dizinleri döndürür."""
    if not os.path.exists(base_path):
        return []
    return [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]

def find_document_files_recursive(directory):
    """Belirtilen dizin ve alt dizinlerindeki tüm PDF ve DOCX dosyalarını bulur."""
    doc_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            # Hem .pdf hem de .docx uzantılı dosyaları kontrol et
            if file.lower().endswith(('.pdf', '.docx')):
                doc_files.append(os.path.join(root, file))
    return doc_files

def test_specific_field(doc_file, field_type):
    """Belirli bir alan tipini test eder ve detaylı çıktı verir."""
    print(f"\n=== {field_type.upper()} TESTİ: {os.path.basename(doc_file)} ===")
    
    try:
        if field_type == "kazanim":
            # Kazanımları test et
            kazanimlar = extract_genel_kazanimlar(doc_file)
            print(f"Bulunan Kazanım Sayısı: {len(kazanimlar)}")
            print("\nBulunan Kazanımlar:")
            for i, kazanim in enumerate(kazanimlar, 1):
                print(f"{i}. {kazanim}")
            return {"field_type": "kazanim", "count": len(kazanimlar), "items": kazanimlar}
            
        elif field_type == "ortamdonanim":
            # Ortam ve donanımı test et
            ortam_donanimi = extract_ortam_donanimi(doc_file)
            print(f"Bulunan Ortam/Donanım Sayısı: {len(ortam_donanimi)}")
            print("\nBulunan Ortam/Donanım Maddeleri:")
            for i, item in enumerate(ortam_donanimi, 1):
                print(f"{i}. {item}")
            
            # Ham metni de göster
            processor = DocumentProcessor()
            tables = processor.extractors['DBF'].get_tables(doc_file)
            raw_content = find_raw_content_for_ortam(tables)
            print(f"\nHam İçerik Arama Sonucu:")
            if raw_content:
                print(f"✅ Bulundu: '{raw_content}'")
            else:
                print("❌ Ham içerik bulunamadı - tabloları kontrol ediliyor...")
                print(f"Toplam tablo sayısı: {len(tables)}")
                for i, table in enumerate(tables):
                    print(f"Tablo {i}: {len(table)} satır")
                    if len(table) > 5 and len(table[5]) > 1:
                        print(f"  Row 6, Col 2: '{table[5][1] if table[5][1] else 'None'}'")
            
            return {"field_type": "ortamdonanim", "count": len(ortam_donanimi), "items": ortam_donanimi, "raw_content": raw_content}
            
        elif field_type == "olcme":
            # Ölçme değerlendirmeyi test et
            olcme_degerlendirme = extract_olcme_degerlendirme(doc_file)
            print(f"Bulunan Ölçme/Değerlendirme Sayısı: {len(olcme_degerlendirme)}")
            print("\nBulunan Ölçme/Değerlendirme Maddeleri:")
            for i, item in enumerate(olcme_degerlendirme, 1):
                print(f"{i}. {item}")
            
            # Ham metni de göster
            raw_content = find_raw_content_for_olcme(doc_file)
            print(f"\nHam İçerik Arama Sonucu:")
            if raw_content:
                print(f"✅ Bulundu: '{raw_content}'")
            else:
                print("❌ Ham içerik bulunamadı - dosya türü kontrol ediliyor...")
                file_ext = doc_file.lower().split('.')[-1]
                print(f"Dosya türü: {file_ext}")
                if file_ext == 'pdf':
                    print("PDF tabloları kontrol ediliyor...")
                elif file_ext == 'docx':
                    print("DOCX tabloları kontrol ediliyor...")
            
            return {"field_type": "olcme", "count": len(olcme_degerlendirme), "items": olcme_degerlendirme, "raw_content": raw_content}
            
        else:
            print(f"Bilinmeyen alan tipi: {field_type}")
            return None
            
    except Exception as e:
        print(f"Hata: {e}")
        return None

def find_raw_content_for_ortam(tables):
    """Ortam ve donanım için ham içeriği bulur."""
    for table_idx, table in enumerate(tables):
        # Önce Row 6, Col 2'yi kontrol et
        if len(table) > 5 and len(table[5]) > 1:
            cell_content = table[5][1] if table[5][1] else ""
            if cell_content and 'Ortam:' in cell_content:
                return cell_content
        
        # Tüm tabloyu tara
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                if cell and isinstance(cell, str) and ('Ortam:' in cell or 'EĞİTİM-ÖĞRETİM ORTAM' in cell.upper()):
                    return cell
    return None

def find_raw_content_for_olcme(file_path):
    """Ölçme değerlendirme için ham içeriği bulur."""
    import pdfplumber
    import docx
    
    debug_info = []
    
    try:
        if file_path.lower().endswith('.docx'):
            debug_info.append("DOCX dosyası işleniyor...")
            doc = docx.Document(file_path)
            debug_info.append(f"Toplam tablo sayısı: {len(doc.tables)}")
            
            for table_idx, table in enumerate(doc.tables):
                debug_info.append(f"Tablo {table_idx}: {len(table.rows)} satır")
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if 'ÖLÇME VE DEĞERLENDİRME' in cell_text.upper():
                            debug_info.append(f"✅ Başlık bulundu - Tablo {table_idx}, Satır {row_idx}, Hücre {cell_idx}")
                            # Sağdaki hücreyi kontrol et
                            if cell_idx + 1 < len(row.cells):
                                content_cell = row.cells[cell_idx + 1].text.strip()
                                if content_cell and 'ÖLÇME VE DEĞERLENDİRME' not in content_cell.upper():
                                    debug_info.append(f"✅ İçerik bulundu sağ hücrede")
                                    print("DEBUG:", "\n".join(debug_info))
                                    return content_cell
                            debug_info.append(f"✅ İçerik aynı hücrede döndürülüyor")
                            print("DEBUG:", "\n".join(debug_info))
                            return cell_text
        else:
            # PDF için - TÜM SAYFALARI KONTROL ET
            with pdfplumber.open(file_path) as pdf:
                debug_info.append(f"PDF dosyası işleniyor - Toplam sayfa: {len(pdf.pages)}")
                
                # Önce tabloları kontrol et
                for page_idx, page in enumerate(pdf.pages):
                    all_tables = page.extract_tables()
                    debug_info.append(f"Sayfa {page_idx}: {len(all_tables)} tablo")
                    
                    for table_idx, table in enumerate(all_tables):
                        for i, row in enumerate(table):
                            for j, cell in enumerate(row):
                                if cell and isinstance(cell, str) and 'ÖLÇME VE DEĞERLENDİRME' in cell.upper():
                                    debug_info.append(f"✅ Başlık bulundu - Sayfa {page_idx}, Tablo {table_idx}, Satır {i}, Hücre {j}")
                                    if ':' in cell:
                                        debug_info.append(f"✅ İçerik aynı hücrede (:) - '{cell}'")
                                        print("DEBUG:", "\n".join(debug_info))
                                        return cell
                                    elif j + 1 < len(row) and row[j + 1]:
                                        content = str(row[j + 1]).strip()
                                        debug_info.append(f"✅ İçerik sağ hücrede - '{content}'")
                                        print("DEBUG:", "\n".join(debug_info))
                                        return content
                                    elif i + 1 < len(table) and len(table[i + 1]) > 0:
                                        content = str(table[i + 1][0]).strip()
                                        debug_info.append(f"✅ İçerik alt satırda - '{content}'")
                                        print("DEBUG:", "\n".join(debug_info))
                                        return content
                
                # Düz metinde ara
                debug_info.append("Tablolarda bulunamadı, düz metinde aranıyor...")
                for page_idx, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and 'ÖLÇME' in text.upper() and 'DEĞERLENDİRME' in text.upper():
                        debug_info.append(f"✅ Düz metinde bulundu - Sayfa {page_idx}")
                        lines = text.split('\n')
                        for i, line in enumerate(lines):
                            if 'ÖLÇME' in line.upper() and 'DEĞERLENDİRME' in line.upper():
                                context_lines = []
                                for j in range(max(0, i-1), min(len(lines), i+5)):
                                    context_lines.append(lines[j])
                                content = ' '.join(context_lines)
                                debug_info.append(f"✅ Bağlam metni oluşturuldu")
                                print("DEBUG:", "\n".join(debug_info))
                                return content
        
        # Hiçbir yerde bulunamadı
        debug_info.append("❌ Hiçbir yerde 'ÖLÇME VE DEĞERLENDİRME' metni bulunamadı")
        print("DEBUG:", "\n".join(debug_info))
        
    except Exception as e:
        debug_info.append(f"❌ Hata: {e}")
        print("DEBUG:", "\n".join(debug_info))
    
    return None

if __name__ == "__main__":
    dbf_base_path = "data/dbf"

    if not os.path.exists(dbf_base_path):
        print(f"Hata: '{dbf_base_path}' dizini bulunamadı.")
        sys.exit(1)

    subdirs = get_subdirectories(dbf_base_path)

    if not subdirs:
        print(f"'{dbf_base_path}' altında alt dizin bulunamadı.")
        sys.exit(0)

    # Komut satırı argümanlarını kontrol et
    if len(sys.argv) < 2:
        print("Kullanım:")
        print("  python test_dbf.py [random | <dizin_adı>]                    # Tüm dosyaları işle")
        print("  python test_dbf.py <dizin_adı> kazanim                      # Sadece kazanımları test et")
        print("  python test_dbf.py <dizin_adı> ortamdonanim                 # Sadece ortam/donanımı test et")
        print("  python test_dbf.py <dizin_adı> olcme                        # Sadece ölçme/değerlendirmeyi test et")
        print(f"\nMevcut dizinler: {', '.join(subdirs)}")
        sys.exit(1)

    # Alan test modunu kontrol et
    field_test_mode = None
    if len(sys.argv) > 2:
        field_test_mode = sys.argv[2].lower()
        if field_test_mode not in ['kazanim', 'ortamdonanim', 'olcme']:
            print(f"Hata: Geçersiz alan tipi '{field_test_mode}'")
            print("Geçerli alan tipleri: kazanim, ortamdonanim, olcme")
            sys.exit(1)

    target_dir = None
    if sys.argv[1].lower() == "random":
        target_dir_name = random.choice(subdirs)
        target_dir = os.path.join(dbf_base_path, target_dir_name)
        print(f"Rasgele seçilen dizin: {target_dir_name}")
    else:
        # Kullanıcı belirli bir dizin adı verdiyse
        specified_dir_name = sys.argv[1]
        if specified_dir_name in subdirs:
            target_dir = os.path.join(dbf_base_path, specified_dir_name)
            print(f"Belirtilen dizin: {specified_dir_name}")
        else:
            print(f"Hata: '{specified_dir_name}' dizini '{dbf_base_path}' altında bulunamadı.")
            print(f"Mevcut dizinler: {', '.join(subdirs)}")
            sys.exit(1)

    if target_dir:
        doc_files_to_process = find_document_files_recursive(target_dir)

        if not doc_files_to_process:
            print(f"'{target_dir}' dizininde PDF veya DOCX dosyası bulunamadı.")
        else:
            if field_test_mode:
                # Belirli alan test modu
                print(f"'{target_dir}' dizininde bulunan {len(doc_files_to_process)} dosya için {field_test_mode.upper()} testi yapılıyor...")
                all_results = []
                
                for doc_file in doc_files_to_process:
                    result = test_specific_field(doc_file, field_test_mode)
                    if result:
                        # Dosya bilgilerini ekle
                        result["file_name"] = os.path.basename(doc_file)
                        result["file_path"] = doc_file
                        all_results.append(result)
                        
                        # JSON kaydetme iptal edildi
                    
            else:
                # Normal tüm dosya işleme modu
                print(f"'{target_dir}' dizininde bulunan {len(doc_files_to_process)} dosya işleniyor...")
                for doc_file in doc_files_to_process:
                    print(f"\n--- {os.path.basename(doc_file)} işleniyor ---")
                    try:
                        # oku_dbf fonksiyonu zaten sonuçları ekrana yazdırıyor
                        result = oku_dbf(doc_file)
                        
                        # JSON kaydetme iptal edildi
                        
                    except Exception as e:
                        print(f"Hata: '{os.path.basename(doc_file)}' işlenirken bir hata oluştu: {e}")
    else:
        print("İşlenecek bir dizin seçilemedi.")