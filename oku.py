import pdfplumber
import docx
import pandas as pd
import re
import os
import glob
import json
from datetime import datetime

def get_tables_from_file(file_path):
    """
    PDF veya DOCX dosyasından tablo listesini döndürür.
    PDF için: pdfplumber, DOCX için: python-docx kullanır.
    """
    if file_path.lower().endswith('.pdf'):
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                tables = []
                for page in pdf.pages:
                    tables.extend(page.extract_tables())
                return tables
        except Exception as e:
            print(f"PDF tablo okuma hatası: {e}")
            return []
    elif file_path.lower().endswith('.docx'):
        try:
            import docx
            doc = docx.Document(file_path)
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)
            return tables
        except Exception as e:
            print(f"DOCX tablo okuma hatası: {e}")
            return []
    else:
        print("Desteklenmeyen dosya formatı.")
        return []

def clean_text(text):
    """
    Metindeki gereksiz \n karakterlerini ve fazla boşlukları temizle
    """
    if not text:
        return text
    
    # Birden fazla boşluğu tek boşlukla değiştir
    text = re.sub(r'\s+', ' ', text.strip())
    return text

def extract_ders_adi(file_path):
    """
    PDF veya DOCX dosyasından ders adını çıkar
    """
    try:
        tables = get_tables_from_file(file_path)
        for table in tables:
            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    if cell and isinstance(cell, str) and 'DERSİN ADI' in cell.upper():
                        # İçerik aynı hücrede ise
                        if ':' in cell:
                            content = cell.split(':', 1)[1].strip()
                            if content:
                                return clean_text(content)
                        # Üst satırda ise (bazı dosyalarda kurs adı üst satırda olabilir)
                        if i > 0:
                            prev_row = table[i - 1]
                            for k in range(len(prev_row)):
                                if prev_row[k] and str(prev_row[k]).strip():
                                    content = str(prev_row[k]).strip()
                                    if content and content != 'None' and len(content) > 3:
                                        if not any(keyword in content.upper() for keyword in ['DERSİN', 'ÖĞRENME', 'EĞİTİM', 'ÖLÇME', 'SINIF']):
                                            return clean_text(content)
                        # Satırın tüm hücrelerinde ara (çünkü veri farklı sütunda olabilir)
                        for k in range(len(row)):
                            if k != j and row[k]:
                                content = str(row[k]).strip()
                                if content and content != 'None' and len(content) > 3:
                                    if not any(keyword in content.upper() for keyword in ['DERSİN', 'ÖĞRENME', 'EĞİTİM', 'ÖLÇME']):
                                        return clean_text(content)
                        # Alt satırda ise
                        if i + 1 < len(table):
                            next_row = table[i + 1]
                            for next_cell in next_row:
                                if next_cell:
                                    content = str(next_cell).strip()
                                    if content and not any(keyword in content.upper() for keyword in ['DERSİN', 'ÖĞRENME', 'EĞİTİM', 'ÖLÇME']):
                                        return clean_text(content)
        return None
    except Exception as e:
        print(f"Error extracting ders adı: {str(e)}")
        return None

def extract_ders_sinifi(file_path):
    """
    PDF veya DOCX dosyasından ders sınıfını çıkar ve sayıya çevir
    """
    try:
        tables = get_tables_from_file(file_path)
        for table in tables:
            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    if cell and isinstance(cell, str) and 'DERSİN SINIFI' in cell.upper():
                        # İçerik aynı hücrede ise
                        if ':' in cell:
                            content = cell.split(':', 1)[1].strip()
                        else:
                            # Üst satırda ise (bazı dosyalarda değer üst satırda olabilir)
                            content = None
                            if i > 0:
                                prev_row = table[i - 1]
                                for k in range(len(prev_row)):
                                    if prev_row[k] and str(prev_row[k]).strip():
                                        cell_content = str(prev_row[k]).strip()
                                        if (cell_content and 'DERSİN' not in cell_content.upper() and
                                            len(cell_content) < 50 and re.search(r'\d', cell_content)):
                                            content = cell_content
                                            break
                            # Satırın tüm hücrelerinde ara
                            if not content:
                                for k in range(len(row)):
                                    if k != j and row[k]:
                                        cell_content = str(row[k]).strip()
                                        if cell_content and 'DERSİN' not in cell_content.upper():
                                            content = cell_content
                                            break
                            if not content:
                                continue
                        if content:
                            match = re.search(r'(\d+)', content)
                            if match:
                                return int(match.group(1))
        return None
    except Exception as e:
        print(f"Error extracting ders sınıfı: {str(e)}")
        return None

def extract_ders_saati(file_path):
    """
    PDF veya DOCX dosyasından haftalık ders saatini çıkar
    """
    try:
        tables = get_tables_from_file(file_path)
        for table in tables:
            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    if cell and isinstance(cell, str) and ('SÜRESİ' in cell.upper() or 'HAFTALIK' in cell.upper()):
                        # İçerik aynı hücrede ise
                        if ':' in cell:
                            content = cell.split(':', 1)[1].strip()
                        else:
                            # Üst satırda ise (bazı dosyalarda değer üst satırda olabilir)
                            content = None
                            if i > 0:
                                prev_row = table[i - 1]
                                for k in range(len(prev_row)):
                                    if prev_row[k] and str(prev_row[k]).strip():
                                        cell_content = str(prev_row[k]).strip()
                                        if (cell_content and 'DERSİN' not in cell_content.upper() and
                                            ('HAFTALIK' in cell_content.upper() or 'SAATİ' in cell_content.upper())):
                                            content = cell_content
                                            break
                            # Satırın tüm hücrelerinde ara
                            if not content:
                                for k in range(len(row)):
                                    if k != j and row[k]:
                                        cell_content = str(row[k]).strip()
                                        if cell_content and 'DERSİN' not in cell_content.upper():
                                            content = cell_content
                                            break
                            if not content:
                                continue
                        if content:
                            match = re.search(r'(\d+)', content)
                            if match:
                                return int(match.group(1))
        return None
    except Exception as e:
        print(f"Error extracting ders saati: {str(e)}")
        return None

def extract_dersin_amaci(file_path):
    """
    PDF veya DOCX dosyasından dersin amacını çıkar
    """
    try:
        tables = get_tables_from_file(file_path)
        for table in tables:
            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    if cell and isinstance(cell, str) and 'DERSİN AMACI' in cell.upper():
                        # Başlıktan sonraki satırlarda ilk anlamlı metni bul
                        for down in range(i + 1, len(table)):
                            next_row = table[down]
                            for next_cell in next_row:
                                if next_cell:
                                    content = str(next_cell).strip()
                                    # En az 10 karakter, başlık değil ve "DERSİN" gibi kelimeler içermiyor
                                    if (content and len(content) > 10 and
                                        not any(kw in content.upper() for kw in ['DERSİN', 'SÜRESİ', 'SINIFI', 'ADI', 'KAZANIM', 'ORTAM', 'DONANIM', 'ÖLÇME'])):
                                        return clean_text(content)
                        # Eğer alt satırlarda bulunamazsa, eski mantıkla devam et
                        content = ""
                        if ':' in cell:
                            content = cell.split(':', 1)[1].strip()
                        if not content and i > 0:
                            prev_row = table[i - 1]
                            for k in range(len(prev_row)):
                                if prev_row[k] and str(prev_row[k]).strip():
                                    cell_content = str(prev_row[k]).strip()
                                    if (cell_content and len(cell_content) > 10 and
                                        'DERSİN' not in cell_content.upper() and
                                        'HAFTALIK' not in cell_content.upper() and
                                        'DERS SAATİ' not in cell_content.upper() and
                                        'SAATİ' not in cell_content.upper()):
                                        content = cell_content
                                        break
                        if not content:
                            for k in range(len(row)):
                                if k != j and row[k]:
                                    cell_content = str(row[k]).strip()
                                    if cell_content and len(cell_content) > 10 and 'DERSİN' not in cell_content.upper():
                                        content = cell_content
                                        break
                        if not content and i + 1 < len(table):
                            next_row = table[i + 1]
                            for next_cell in next_row:
                                if next_cell and str(next_cell).strip():
                                    content = str(next_cell).strip()
                                    break
                        if content and len(content) > 10:
                            return clean_text(content)
        return None
    except Exception as e:
        print(f"Error extracting dersin amacı: {str(e)}")
        return None

def extract_genel_kazanimlar(file_path):
    """
    PDF veya DOCX dosyasından dersin genel kazanımlarını çıkar
    """
    try:
        if file_path.lower().endswith('.docx'):
            # DOCX dosyası için
            import docx
            doc = docx.Document(file_path)
            
            kazanimlar = []
            
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if 'DERSİN KAZANIMLARI' in cell_text.upper():
                            # Eğer başlık hücresiyse, sağındaki veya alt hücreyi kontrol et
                            content_cell = None
                            
                            # Sağdaki hücreyi kontrol et
                            if cell_idx + 1 < len(row.cells):
                                content_cell = row.cells[cell_idx + 1].text.strip()
                            
                            # Eğer sağda yoksa, aynı hücrenin içeriğini kontrol et
                            if not content_cell or 'DERSİN KAZANIMLARI' in content_cell.upper():
                                content_cell = cell_text
                            
                            if content_cell:
                                # Kazanımları satır satır ayır
                                lines = content_cell.split('\n')
                                for line in lines:
                                    line = line.strip()
                                    if line and len(line) > 10 and not any(keyword in line.upper() for keyword in ['DERSİN', 'KAZANIM']):
                                        # Numbering'i kaldır
                                        cleaned_line = re.sub(r'^\d+\.\s*', '', line)
                                        if cleaned_line and len(cleaned_line) > 10:
                                            kazanimlar.append(clean_text(cleaned_line))
            
            return kazanimlar[:10]
            
        else:
            # PDF dosyası için (eski kod)
            with pdfplumber.open(file_path) as pdf:
                page = pdf.pages[0]
                tables = page.extract_tables()
                
                kazanimlar = []
                
                for table in tables:
                    for i, row in enumerate(table):
                        for j, cell in enumerate(row):
                            if cell and isinstance(cell, str) and ('DERSİN ÖĞRENME KAZANIMLARI' in cell.upper()):
                                
                                # İçerik aynı hücrede ise
                                if ':' in cell:
                                    content = cell.split(':', 1)[1].strip()
                                else:
                                    # Yan hücredeki içeriği al
                                    if j + 1 < len(row) and row[j + 1]:
                                        content = str(row[j + 1]).strip()
                                    else:
                                        # Alt satırdaki içeriği al
                                        if i + 1 < len(table) and len(table[i + 1]) > j:
                                            content = str(table[i + 1][j]).strip() if table[i + 1][j] else ""
                                        else:
                                            continue
                                
                                if content and len(content) > 10:
                                    # Madde işaretleriyle ayrılmış kazanımları ayır
                                    if '•' in content or content.startswith(('1.', '2.', '3.')):
                                        kazanim_items = re.split(r'[•\-]\s*|\d+\.\s*', content)
                                        for item in kazanim_items:
                                            clean_item = clean_text(item)
                                            if clean_item and len(clean_item) > 10:
                                                kazanimlar.append(clean_item)
                                    else:
                                        clean_item = clean_text(content)
                                        if clean_item and len(clean_item) > 10:
                                            kazanimlar.append(clean_item)
                                
                                return kazanimlar[:10]  # İlk eşleşmede dön
                
                return kazanimlar
        
    except Exception as e:
        print(f"Error extracting genel kazanımlar: {str(e)}")
        return []

def extract_ortam_donanimi(file_path):
    """
    PDF veya DOCX dosyasından eğitim-öğretim ortam ve donanımını çıkar
    Ortam: X ve Y ortamı. Donanım: A, B ve C araçları formatında parse eder
    """
    try:
        if file_path.lower().endswith('.docx'):
            # DOCX dosyası için
            import docx
            doc = docx.Document(file_path)
            
            ortam_donanimi = []
            
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if 'EĞİTİM-ÖĞRETİM ORTAM' in cell_text.upper() or 'ORTAM VE DONANIMI' in cell_text.upper():
                            # Eğer başlık hücresiyse, sağındaki hücreyi kontrol et
                            content_cell = None
                            
                            # Sağdaki hücreyi kontrol et
                            if cell_idx + 1 < len(row.cells):
                                content_cell = row.cells[cell_idx + 1].text.strip()
                            
                            # Eğer sağda yoksa, aynı hücrenin içeriğini kontrol et
                            if not content_cell or ('EĞİTİM-ÖĞRETİM ORTAM' in content_cell.upper() or 'ORTAM VE DONANIMI' in content_cell.upper()):
                                content_cell = cell_text
                            
                            if content_cell:
                                return parse_ortam_donanimi_content(content_cell)
            
            return ortam_donanimi
            
        else:
            # PDF dosyası için (eski kod)
            with pdfplumber.open(file_path) as pdf:
                page = pdf.pages[0]  # İlk sayfa
                tables = page.extract_tables()
                
                ortam_donanimi = []
                
                for table_idx, table in enumerate(tables):
                    
                    # Önce Row 6, Col 2'yi kontrol et (çünkü önceden çalışıyordu)
                    if len(table) > 5 and len(table[5]) > 1:
                        cell_content = table[5][1] if table[5][1] else ""
                        
                        if cell_content and 'Ortam:' in cell_content:
                            
                            # Metni tek satır haline getir
                            cleaned_content = re.sub(r'\s+', ' ', cell_content.strip())
                            
                            # Özel düzeltme - eğer gerekliyse
                            if 'tahta/projeksiyon, çizim masası, çizim seti ile iş ortamı Donanım: Akıllı' in cleaned_content:
                                fixed_content = cleaned_content.replace(
                                    'tahta/projeksiyon, çizim masası, çizim seti ile iş ortamı Donanım: Akıllı',
                                    'Donanım: Akıllı tahta/projeksiyon, çizim masası, çizim seti ile iş ortamı'
                                )
                                return parse_ortam_donanimi_content(fixed_content)
                            else:
                                return parse_ortam_donanimi_content(cleaned_content)
                    
                    # Tüm tabloyu tara
                    for i, row in enumerate(table):
                        for j, cell in enumerate(row):
                            if cell and isinstance(cell, str) and ('Ortam:' in cell or 'EĞİTİM-ÖĞRETİM ORTAM' in cell.upper()):
                                content = cell
                                
                                # Metni tek satır haline getir
                                cleaned_content = re.sub(r'\s+', ' ', content.strip())
                                
                                if len(cleaned_content) > 10:
                                    return parse_ortam_donanimi_content(cleaned_content)
                
                return ortam_donanimi
        
    except Exception as e:
        print(f"Error extracting ortam donanımı: {str(e)}")
        return []

def parse_ortam_donanimi_content(content):
    """
    Ortam ve donanım içeriğini maddelere ayır
    """
    ortam_donanimi = []
    
    # Ortam ve Donanım bölümlerini ayır
    ortam_pattern = r'Ortam:\s*([^.]*?)\.?\s*(?:Donanım:|$)'
    donanim_pattern = r'Donanım:\s*(.+?)(?:sağlanmalıdır|$)'
    
    ortam_match = re.search(ortam_pattern, content, re.IGNORECASE | re.DOTALL)
    donanim_match = re.search(donanim_pattern, content, re.IGNORECASE | re.DOTALL)
    
    # Ortam kısmını parse et
    if ortam_match:
        ortam_text = ortam_match.group(1).strip()
        
        # Virgül ve "ve" ile ayrılmış ortamları ayır
        ortam_items = re.split(r',\s*|\s+ve\s+', ortam_text)
        for item in ortam_items:
            clean_item = clean_text(item.strip().rstrip('.,'))
            if clean_item and len(clean_item) > 3:
                ortam_donanimi.append(clean_item.title())
    
    # Donanım kısmını parse et  
    if donanim_match:
        donanim_text = donanim_match.group(1).strip()
        
        # "sağlanmalıdır" kısmını temizle
        donanim_text = re.sub(r'\s*(ile\s+)?iş\s+ortam[ıi]\s+sağlanmalıdır.*$', '', donanim_text, flags=re.IGNORECASE)
        
        # Virgül ile ayrılmış donanımları ayır
        donanim_items = re.split(r',\s*', donanim_text)
        for item in donanim_items:
            # Son maddedeki "ve" yi temizle
            if ' ve ' in item:
                ve_parts = re.split(r'\s+ve\s+', item)
                for part in ve_parts:
                    clean_item = clean_text(part.strip().rstrip('.,'))
                    if clean_item and len(clean_item) > 3:
                        ortam_donanimi.append(clean_item.title())
            else:
                clean_item = clean_text(item.strip().rstrip('.,'))
                if clean_item and len(clean_item) > 3:
                    ortam_donanimi.append(clean_item.title())
    
    # Eğer pattern match etmezse basit ayrıştırma yap
    if not ortam_match and not donanim_match:
        # Tüm noktalama ile ayrılmış öğeleri al
        items = re.split(r'[,.;]\s*|\s+ve\s+', content)
        for item in items:
            clean_item = clean_text(item.strip().rstrip('.,'))
            if clean_item and len(clean_item) > 3 and 'sağlanmalıdır' not in clean_item.lower():
                ortam_donanimi.append(clean_item.title())
    
    return ortam_donanimi

def extract_kazanim_tablosu(file_path):
    """
    PDF veya DOCX dosyasından kazanım sayısı ve süre tablosunu çıkar
    """
    try:
        if file_path.lower().endswith('.docx'):
            # DOCX dosyası için
            import docx
            doc = docx.Document(file_path)
            
            kazanim_tablosu = []
            
            for table in doc.tables:
                # Kazanım tablosu başlığını bul
                found_header = False
                data_start_row = None
                
                for row_idx, row in enumerate(table.rows):
                    row_cells = [cell.text.strip() for cell in row.cells]
                    row_text = ' '.join(row_cells).upper()
                    
                    # Header row'u bul: "ÖĞRENME BİRİMİ/ÜNİTE", "KAZANIM SAYISI", "DERS SAATİ"
                    if ('ÖĞRENME BİRİMİ' in row_text or 'ÜNİTE' in row_text) and 'KAZANIM SAYISI' in row_text and ('DERS SAATİ' in row_text or 'SAAT' in row_text):
                        found_header = True
                        data_start_row = row_idx + 1
                        break
                
                # Eğer header bulunduysa, veri satırlarını işle
                if found_header and data_start_row is not None:
                    for row_idx in range(data_start_row, len(table.rows)):
                        row = table.rows[row_idx]
                        row_cells = [cell.text.strip() for cell in row.cells]
                        
                        # Boş satırları atla
                        if all(cell == '' for cell in row_cells):
                            continue
                        
                        # TOPLAM satırını atla
                        row_text = ' '.join(row_cells).upper()
                        if 'TOPLAM' in row_text:
                            continue
                        
                        # En az 4 sütun olmalı: [başlık], [birim adı], [kazanım sayısı], [ders saati], [oran]
                        if len(row_cells) >= 4:
                            # İkinci sütun birim adı olmalı
                            birim_adi = clean_text(row_cells[1]) if row_cells[1] else ""
                            kazanim_sayisi = clean_text(row_cells[2]) if row_cells[2] else ""
                            ders_saati = clean_text(row_cells[3]) if row_cells[3] else ""
                            
                            # Geçerli veri kontrolü
                            if (birim_adi and len(birim_adi) > 3 and 
                                kazanim_sayisi.isdigit() and 
                                ders_saati.replace(',', '.').replace('.', '').isdigit()):
                                
                                kazanim_tablosu.append({
                                    "birim_adi": birim_adi,
                                    "kazanim_sayisi": kazanim_sayisi,
                                    "ders_saati_orani": ders_saati
                                })
                    
                    # Eğer veri bulunmuşsa bu tablodan çık
                    if kazanim_tablosu:
                        break
            
            return kazanim_tablosu
            
        else:
            # PDF dosyası için (eski kod)
            with pdfplumber.open(file_path) as pdf:
                kazanim_tablosu = []
                
                # Özel olarak Page 2'deki first table'ı kontrol et (sayısal veriler içeren)
                if len(pdf.pages) > 1:
                    page = pdf.pages[1]  # Page 2
                    tables = page.extract_tables()
                    
                    if tables:
                        table = tables[0]  # İlk tablo (sayısal veriler)
                        
                        for row_idx, row in enumerate(table):
                            if row and len(row) >= 4:
                                # İlk sütun boş veya None değilse ve ikinci sütunda öğrenme birimi varsa
                                if row[1] and str(row[1]).strip() and str(row[1]).strip() != 'None':
                                    birim_adi = clean_text(str(row[1]))
                                    kazanim_sayisi = clean_text(str(row[2])) if row[2] else ""
                                    ders_saati = clean_text(str(row[3])) if row[3] else ""
                                    
                                    # TOPLAM satırını atla ve geçerli öğrenme birimi kontrolü
                                    if ('TOPLAM' not in birim_adi.upper() and 
                                        len(birim_adi) > 3 and
                                        kazanim_sayisi.isdigit() and 
                                        ders_saati.replace(',', '.').replace('.', '').isdigit()):
                                        
                                        kazanim_tablosu.append({
                                            "birim_adi": birim_adi,
                                            "kazanim_sayisi": kazanim_sayisi,
                                            "ders_saati_orani": ders_saati
                                        })
                
                # Eğer Page 2'de bulamazsa, eski yöntemi dene (metin tabanlı arama)
                if not kazanim_tablosu:
                    for page_num, page in enumerate(pdf.pages):
                        tables = page.extract_tables()
                        
                        for table in tables:
                            # "KAZANIM SAYISI VE SÜRE TABLOSU" başlığını bul
                            found_table = False
                            table_start_row = None
                            
                            for i, row in enumerate(table):
                                row_text = ' '.join([str(cell) if cell else '' for cell in row]).upper()
                                if 'KAZANIM SAYISI' in row_text and ('SÜRE' in row_text or 'TABLOSU' in row_text):
                                    found_table = True
                                    table_start_row = i
                                    break
                            
                            if found_table and table_start_row is not None:
                                # Tablo başlık satırlarını bul
                                header_found = False
                                for i in range(table_start_row, len(table)):
                                    row = table[i]
                                    row_text = ' '.join([str(cell) if cell else '' for cell in row]).upper()
                                    
                                    # Sütun başlıklarını bul
                                    if any(keyword in row_text for keyword in ['ÖĞRENME BİRİMİ', 'KAZANIM', 'SÜRE', 'DERS SAATİ']):
                                        header_found = True
                                        # Veri satırlarını oku
                                        for k in range(i + 1, len(table)):
                                            data_row = table[k]
                                            
                                            # Boş satırları atla
                                            row_values = [str(cell).strip() if cell else '' for cell in data_row]
                                            if all(val in ['', 'None'] for val in row_values):
                                                continue
                                            
                                            # TOPLAM satırını atla
                                            row_text_data = ' '.join(row_values).upper()
                                            if 'TOPLAM' in row_text_data:
                                                continue
                                            
                                            # Geçerli veri olan satırları al
                                            valid_data = [val for val in row_values if val not in ['', 'None']]
                                            
                                            if len(valid_data) >= 2:
                                                birim_adi = clean_text(valid_data[0])
                                                kazanim_sayisi = clean_text(valid_data[1]) if len(valid_data) > 1 else ""
                                                ders_saati = clean_text(valid_data[2]) if len(valid_data) > 2 else ""
                                                
                                                # Birim adının geçerli olduğunu kontrol et
                                                if birim_adi and len(birim_adi) > 3 and not birim_adi.isdigit():
                                                    kazanim_tablosu.append({
                                                        "birim_adi": birim_adi,
                                                        "kazanim_sayisi": kazanim_sayisi,
                                                        "ders_saati_orani": ders_saati
                                                    })
                                        break
                                
                                if header_found:
                                    return kazanim_tablosu
                
                return kazanim_tablosu
        
    except Exception as e:
        print(f"Error extracting kazanım tablosu: {str(e)}")
        return []

def extract_ogrenme_birimleri(file_path):
    """
    PDF veya DOCX dosyasından öğrenme birimlerini detaylı olarak çıkar
    3 sütunlu tablo: ÖĞRENME BİRİMİ | KONULAR | KAZANIMLAR
    """
    try:
        if file_path.lower().endswith('.docx'):
            # DOCX dosyası için
            import docx
            doc = docx.Document(file_path)
            
            ogrenme_birimleri = []
            
            for table in doc.tables:
                # 3 sütunlu öğrenme birimleri tablosunu bul
                if table.rows:
                    first_row = table.rows[0]
                    first_row_cells = [cell.text.strip() for cell in first_row.cells]
                    row_text = ' '.join(first_row_cells).upper()
                    
                    # Header kontrolü: ÖĞRENME BİRİMİ, KONULAR, KAZANIMLARI
                    if ('ÖĞRENME BİRİMİ' in row_text and 'KONULAR' in row_text and 
                        'KAZANIMLARI' in row_text):
                        
                        # Veri satırlarını işle (header'dan sonraki satırlar)
                        for row_idx in range(1, len(table.rows)):
                            row = table.rows[row_idx]
                            row_cells = [cell.text.strip() for cell in row.cells]
                            
                            if len(row_cells) >= 3:
                                birim_adi = clean_text(row_cells[0]) if row_cells[0] else ""
                                konular_text = row_cells[1] if row_cells[1] else ""
                                kazanimlar_text = row_cells[2] if row_cells[2] else ""
                                
                                if birim_adi and len(birim_adi) > 3:
                                    # Konuları parse et
                                    konular = []
                                    if konular_text:
                                        # Satır satır ayır ve temizle
                                        konu_lines = konular_text.split('\n')
                                        for line in konu_lines:
                                            line = line.strip()
                                            if line and len(line) > 5:
                                                # Numaralandırma varsa kaldır
                                                clean_line = re.sub(r'^\d+\.\s*', '', line)
                                                if clean_line:
                                                    konular.append(clean_text(clean_line))
                                    
                                    # Ana kazanımları parse et
                                    kazanimlar = []
                                    if kazanimlar_text:
                                        # Önce numaralı kazanımları dene
                                        ana_kazanim_pattern = r'(\d+\.\s*[^•\uf0b7\n]*?)(?=\n[\s]*[\uf0b7•]|\n\d+\.|\Z)'
                                        ana_kazanimlar = re.findall(ana_kazanim_pattern, kazanimlar_text, re.DOTALL)
                                        
                                        if ana_kazanimlar:
                                            # Numaralı kazanımlar bulundu
                                            for kazanim in ana_kazanimlar:
                                                kazanim = kazanim.strip()
                                                if kazanim and len(kazanim) > 10:
                                                    # Numarayı kaldır
                                                    clean_kazanim = re.sub(r'^\d+\.\s*', '', kazanim)
                                                    if clean_kazanim:
                                                        kazanimlar.append(clean_text(clean_kazanim))
                                        else:
                                            # Numaralı kazanım yoksa, cümle sonlarını al
                                            lines = kazanimlar_text.split('\n')
                                            for line in lines:
                                                line = line.strip()
                                                if (line and len(line) > 10 and 
                                                    (line.endswith('.') or line.endswith('r.')) and
                                                    not any(char in line for char in ['•', '\uf0b7'])):
                                                    kazanimlar.append(clean_text(line))
                                    
                                    birim = {
                                        "birim_adi": birim_adi,
                                        "konular": konular,
                                        "kazanimlar": kazanimlar
                                    }
                                    
                                    ogrenme_birimleri.append(birim)
                        
                        # Tablo bulunduysa döngüden çık
                        if ogrenme_birimleri:
                            break
            
            return ogrenme_birimleri
            
        else:
            # PDF dosyası için (geliştirilmiş kod)
            with pdfplumber.open(file_path) as pdf:
                all_tables = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    all_tables.extend(tables)
            
            ogrenme_birimleri = []
            
            # Öğrenme birimlerini içeren tabloları bul (basitleştirilmiş yaklaşım)
            # Analiz sonucuna göre tables 5, 6, 7'de öğrenme birimleri var
            learning_unit_tables = []
            
            # Önce header'ı olan ana tabloyu bul
            for table_idx, table in enumerate(all_tables):
                if not table or len(table) < 2:
                    continue
                
                for row in table:
                    if row and len(row) > 0:
                        row_text = ' '.join([str(cell) if cell else '' for cell in row]).upper()
                        if ('ÖĞRENME BİRİMİ' in row_text and 'KONULAR' in row_text):
                            learning_unit_tables.append(table_idx)
                            # Bu tablodan sonraki 2-3 tabloyu da ekle (devam tabloları)
                            for next_idx in range(table_idx + 1, min(table_idx + 4, len(all_tables))):
                                if next_idx < len(all_tables):
                                    learning_unit_tables.append(next_idx)
                            break
                if learning_unit_tables:
                    break
            
            # Eğer header tablosu bulunamazsa, bilinen tablo indekslerini kullan
            if not learning_unit_tables and len(all_tables) > 7:
                learning_unit_tables = [5, 6, 7]
            
            # Öğrenme birimlerini çıkar
            seen_units = set()  # Dublicate kontrolü
            
            for table_idx in learning_unit_tables:
                table = all_tables[table_idx]
                
                for row_idx, row in enumerate(table):
                    if not row or len(row) < 2:
                        continue
                    
                    birim_adi = str(row[0]).strip() if row[0] else ""
                    if not birim_adi or len(birim_adi) < 5:
                        continue
                    
                    # Header satırlarını atla
                    birim_adi_upper = birim_adi.upper()
                    if any(header in birim_adi_upper for header in ['ÖĞRENME BİRİMİ', 'KONULAR', 'KAZANIM', 'AÇIKLAMA']):
                        continue
                    
                    # Öğrenme birimi ismini temizle
                    birim_adi = birim_adi.replace('\n', ' ').strip()
                    
                    # Öğrenme birimi kontrolü - daha spesifik
                    expected_units = [
                        "SAYI SİSTEMLERİ", "DATA ÇEVİRİCİLER", 
                        "DİSPLAYLER", "KOKPİT ALETLER",
                        "ELEKTROSTATİK DEŞARJ", "ELEKTROMANYETİK",
                        "DİJİTAL UÇAK SİSTEMLERİ",
                        "FİBER OPTİK",
                        "KABİN BAKIM"
                    ]
                    
                    is_learning_unit = False
                    
                    # Tam eşleşme kontrolü
                    for expected in expected_units:
                        expected_words = expected.split()
                        matches = sum(1 for word in expected_words if word in birim_adi_upper)
                        if matches >= len(expected_words) // 2:  # En az yarısı eşleşsin
                            is_learning_unit = True
                            break
                    
                    # Bilinen öğrenme birimi isimlerini kontrol et
                    if not is_learning_unit:
                        known_units = [
                            "SAYI SİSTEMLERİ VE DATA ÇEVİRİCİLER",
                            "DİSPLAYLER VE KOKPİT ALETLERİ",
                            "ELEKTROSTATİK DEŞARJ VE ELEKTROMANYETİK ÇEVRE",
                            "DİJİTAL UÇAK SİSTEMLERİ",
                            "FİBER OPTİK",
                            "KABİN BAKIM"
                        ]
                        
                        for known in known_units:
                            known_words = known.split()
                            matches = sum(1 for word in known_words if word in birim_adi_upper)
                            if matches >= 2:  # En az 2 kelime eşleşsin
                                is_learning_unit = True
                                break
                    
                    if not is_learning_unit:
                        continue
                    
                    # Dublicate kontrolü
                    if birim_adi in seen_units:
                        continue
                    seen_units.add(birim_adi)
                    
                    # Konuları çıkar (sütun 1)
                    konular = []
                    if len(row) > 1 and row[1]:
                        topics_text = str(row[1])
                        # Numaralı konuları bul
                        topic_matches = re.findall(r'(\d+\.\s*[^.\n\d]+)', topics_text)
                        if topic_matches:
                            for topic in topic_matches:
                                clean_topic = re.sub(r'^\d+\.\s*', '', topic.strip())
                                if clean_topic and len(clean_topic) > 5:
                                    konular.append(clean_text(clean_topic))
                        else:
                            # Satır satır ayır
                            lines = topics_text.split('\n')
                            for line in lines:
                                line = line.strip()
                                if line and len(line) > 5:
                                    clean_line = re.sub(r'^\d+\.\s*', '', line)
                                    if clean_line and not clean_line.startswith('•'):
                                        konular.append(clean_text(clean_line))
                    
                    # Kazanımları çıkar (sütun 2)
                    kazanimlar = []
                    achievements_text = ""
                    
                    if len(row) > 2 and row[2]:
                        achievements_text = str(row[2])
                    elif len(row) > 3 and row[3]:  # Bazen 4. sütunda olabiliyor
                        achievements_text = str(row[3])
                    
                    if achievements_text:
                        # Ana kazanımları bul (numara ile başlayanlar)
                        achievement_matches = re.findall(r'(\d+\.\s*[^•\uf0b7\n]*?)(?=\n[\s]*[\uf0b7•]|\n\d+\.|\Z)', achievements_text, re.DOTALL)
                        if achievement_matches:
                            for achievement in achievement_matches:
                                achievement = achievement.strip()
                                if achievement and len(achievement) > 10:
                                    clean_achievement = re.sub(r'^\d+\.\s*', '', achievement)
                                    if clean_achievement:
                                        kazanimlar.append(clean_text(clean_achievement))
                        else:
                            # Cümle sonlarına göre ayır
                            sentences = re.findall(r'[^.]+\.', achievements_text)
                            for sentence in sentences:
                                sentence = sentence.strip()
                                if sentence and len(sentence) > 15 and not sentence.startswith('•'):
                                    kazanimlar.append(clean_text(sentence))
                    
                    # Devam eden tablolarda ek veri kontrolü
                    if table_idx < len(all_tables) - 1:
                        next_table = all_tables[table_idx + 1]
                        if next_table and len(next_table) > 0:
                            first_row_next = next_table[0]
                            if (first_row_next and len(first_row_next) > 2 and 
                                first_row_next[0] is None and first_row_next[2]):
                                continuation_text = str(first_row_next[2])
                                if continuation_text and len(continuation_text) > 20:
                                    additional_achievements = re.findall(r'[^.]+\.', continuation_text)
                                    for sentence in additional_achievements:
                                        sentence = sentence.strip()
                                        if sentence and len(sentence) > 15:
                                            kazanimlar.append(clean_text(sentence))
                    
                    if birim_adi and (konular or kazanimlar):
                        birim = {
                            "birim_adi": clean_text(birim_adi),
                            "konular": konular,
                            "kazanimlar": kazanimlar
                        }
                        ogrenme_birimleri.append(birim)
            
            return ogrenme_birimleri
        
    except Exception as e:
        print(f"Error extracting öğrenme birimleri: {str(e)}")
        return []

def extract_uygulama_faaliyetleri(pdf_path):
    """
    PDF'den uygulama faaliyetlerini çıkar
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            all_tables = []
            for page in pdf.pages:
                tables = page.extract_tables()
                all_tables.extend(tables)
        
        uygulama_faaliyetleri = []
        
        for table in all_tables:
            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    if cell and isinstance(cell, str) and ('UYGULAMA FAALİYETLERİ' in cell.upper() or 'TEMRİNLER' in cell.upper()):
                        # Alt satırlardaki faaliyetleri topla
                        for k in range(i + 1, len(table)):
                            next_row = table[k]
                            for next_cell in next_row:
                                if next_cell:
                                    content = str(next_cell).strip()
                                    if content and content != 'None' and len(content) > 10:
                                        # Madde işaretleriyle ayrılmış faaliyetleri ayır
                                        if '•' in content or '-' in content or content.startswith(('1.', '2.', '3.')):
                                            faaliyet_items = re.split(r'[•\-]\s*|\d+\.\s*', content)
                                            for item in faaliyet_items:
                                                clean_item = clean_text(item)
                                                if clean_item and len(clean_item) > 10:
                                                    uygulama_faaliyetleri.append(clean_item)
                                        else:
                                            clean_item = clean_text(content)
                                            if clean_item and len(clean_item) > 10:
                                                uygulama_faaliyetleri.append(clean_item)
        
        return uygulama_faaliyetleri[:10]  # İlk 10 faaliyet
        
    except Exception as e:
        print(f"Error extracting uygulama faaliyetleri: {str(e)}")
        return []

def extract_olcme_degerlendirme(file_path):
    """
    PDF veya DOCX dosyasından ölçme ve değerlendirme yöntemlerini çıkar
    Basit regex ile ölçme araçlarını tespit eder
    """
    try:
        if file_path.lower().endswith('.docx'):
            # DOCX dosyası için
            import docx
            doc = docx.Document(file_path)
            
            olcme_yontemleri = []
            
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if 'ÖLÇME VE DEĞERLENDİRME' in cell_text.upper():
                            # Eğer başlık hücresiyse, sağındaki hücreyi kontrol et
                            content_cell = None
                            
                            # Sağdaki hücreyi kontrol et
                            if cell_idx + 1 < len(row.cells):
                                content_cell = row.cells[cell_idx + 1].text.strip()
                            
                            # Eğer sağda yoksa, aynı hücrenin içeriğini kontrol et
                            if not content_cell or 'ÖLÇME VE DEĞERLENDİRME' in content_cell.upper():
                                content_cell = cell_text
                            
                            if content_cell:
                                olcme_yontemleri.extend(analyze_measurement_content(content_cell))
            
            # Dublicate'leri temizle
            return list(set(olcme_yontemleri))
            
        else:
            # PDF dosyası için (eski kod)
            with pdfplumber.open(file_path) as pdf:
                olcme_yontemleri = []
                
                # Önce tablolarda ara (mevcut kod)
                page = pdf.pages[0]
                all_tables = page.extract_tables()
                
                for table in all_tables:
                    for i, row in enumerate(table):
                        for j, cell in enumerate(row):
                            if cell and isinstance(cell, str) and 'ÖLÇME VE DEĞERLENDİRME' in cell.upper():
                                # İçerik aynı hücrede ise
                                if ':' in cell:
                                    content = cell.split(':', 1)[1].strip()
                                else:
                                    # Yan hücrede ise
                                    if j + 1 < len(row) and row[j + 1]:
                                        content = str(row[j + 1]).strip()
                                    else:
                                        # Alt satırda ise
                                        if i + 1 < len(table) and len(table[i + 1]) > 0:
                                            next_row = table[i + 1]
                                            content = str(next_row[0]).strip() if next_row[0] else ""
                                        else:
                                            continue
                                
                                if content and content != 'None':
                                    olcme_yontemleri.extend(analyze_measurement_content(content))
                
                # Eğer tablolarda bulamadıysa, düz metinde ara (yeni eklenen)
                if not olcme_yontemleri:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            lines = text.split('\n')
                            for i, line in enumerate(lines):
                                # Çok satırlı "ÖLÇME VE DEĞERLENDİRME" başlığını kontrol et
                                if ('ÖLÇME' in line.upper() and 'VE' in line.upper()) or 'DEĞERLENDİRME' in line.upper():
                                    # Bu satır ve sonraki birkaç satırı birleştir
                                    context_lines = []
                                    for j in range(max(0, i-1), min(len(lines), i+5)):
                                        context_lines.append(lines[j])
                                    
                                    full_context = ' '.join(context_lines)
                                    if 'ÖLÇME' in full_context.upper() and 'DEĞERLENDİRME' in full_context.upper():
                                        olcme_yontemleri.extend(analyze_measurement_content(full_context))
                                        if olcme_yontemleri:  # Bulunca çık
                                            break
                            if olcme_yontemleri:  # Bulunca çık
                                break
            
            # Dublicate'leri temizle
            return list(set(olcme_yontemleri))
        
    except Exception as e:
        print(f"Error extracting ölçme değerlendirme: {str(e)}")
        return []

def analyze_measurement_content(content):
    """
    Verilen metinde ölçme araçlarını tespit et
    """
    olcme_yontemleri = []
    
    # Bilinen ölçme araçlarını tespit et (yazım varyasyonları dahil)
    olcme_araclari = [
        'gözlem formu', 'gozlem formu', 'derecelendirme ölçeği', 'derecelendirme olceği',
        'dereceli puanlama anahtarı', 'dereceli puanlama anahtari',
        'öz değerlendirme', 'oz degerlendirme', 'öz-değerlendirme',
        'akran değerlendirme', 'akran degerlendirme', 'akran değerlendirmesi',
        'performans değerlendirme', 'performans degerlendirme', 'performans değerlendirmesi',
        'portfolyo', 'portföy', 'portfolio', 'proje değerlendirme', 'proje degerlendirme',
        'rubrik', 'rubric', 'kontrol listesi', 'açık uçlu', 'acik uclu', 'açık-uçlu',
        'çoktan seçmeli', 'coktan secmeli', 'çoktan-seçmeli', 'doğru yanlış', 'dogru yanlis',
        'doğru-yanlış', 'eşleştirme', 'eslestirme', 'geçerlik', 'gecerlik',
        'güvenirlik', 'guvenirlik', 'güvenilirlik', 'biçimlendirici değerlendirme',
        'bicimlendirici degerlendirme', 'düzey belirleyici değerlendirme',
        'duzey belirleyici degerlendirme', 'tanılayıcı değerlendirme', 'tanilayici degerlendirme',
        'otantik değerlendirme', 'otantik degerlendirme', 'özgün değerlendirme',
        'performansa dayalı değerlendirme', 'performansa dayali degerlendirme',
        'e-sınav', 'e sinav', 'elektronik sınav', 'elektronik sinav',
        'e-sınav merkezi', 'e sinav merkezi', 'anket', 'açık uçlu soru', 'acik uclu soru',
        'çoktan seçmeli test', 'coktan secmeli test', 'klasik sınav', 'klasik sinav',
        'kısa cevaplı soru', 'kisa cevapli soru', 'doğru–yanlış testi', 'dogru-yanlis testi',
        'eşleştirme testi', 'eslestirme testi', 'ölçme aracı', 'olcme araci',
        'puanlama anahtarı', 'puanlama anahtari', 'çeldirici', 'celdirici',
        'madde analizi', 'deneme uygulaması', 'deneme uygulamasi',
        'bireyselleştirilmiş eğitim programı', 'bireysellestirilmis egitim programi',
        'bep', 'b.e.p', 'öğrenci ürün dosyası', 'ogrenci urun dosyasi',
        'performans görevi', 'performans gorevi', 'proje ödevi', 'proje odevi',
        'izleme değerlendirme', 'izleme degerlendirme', 'süreç değerlendirme',
        'surec degerlendirme', 'ürün değerlendirme', 'urun degerlendirme',
        'otomatik puanlama', 'elektronik ölçme araçları', 'elektronik olcme araclari',
        'karne notu', 'başarı puanı', 'basari puani', 'sınav uygulama kılavuzu',
        'sinav uygulama kilavuzu', 'değerlendirme kriteri', 'degerlendirme kriteri',
        'geri bildirim', 'geribildirim', 'puanlama ölçeği', 'puanlama olcegi',
        'öğrenme çıktısı', 'ogrenme ciktisi', 'zümre değerlendirmesi', 'zumre degerlendirmesi',
        'öğretmen gözlemi', 'ogretmen gozlemi', 'öğrenci öz değerlendirmesi',
        'ogrenci oz degerlendirmesi', 'ölçme ve değerlendirme uzmanı',
        'olcme ve degerlendirme uzmani', 'ölçme değerlendirme merkezi',
        'olcme degerlendirme merkezi', 'veri analizi', 'raporlama',
        'değerlendirme formu', 'degerlendirme formu'
    ]
    
    content_lower = content.lower()
    for arac in olcme_araclari:
        if arac in content_lower:
            # Formu/formları kelimelerini temizle
            clean_arac = arac.replace(' formu', '').replace(' formları', '')
            if clean_arac not in olcme_yontemleri:
                olcme_yontemleri.append(clean_arac.title())
    
    return olcme_yontemleri

def oku(file_path):
    """
    PDF veya DOCX dosyasından tüm ders bilgilerini çıkar ve JSON formatında döndür
    """
    try:
        filename = os.path.basename(file_path)
        
        # Temel bilgileri çıkar
        ders_adi = extract_ders_adi(file_path)
        ders_sinifi = extract_ders_sinifi(file_path)
        ders_saati = extract_ders_saati(file_path)
        dersin_amaci = extract_dersin_amaci(file_path)
        # Diğer fonksiyonlar henüz güncellenmediği için eski parametreyle devam ediyor
        genel_kazanimlar = extract_genel_kazanimlar(file_path)
        ortam_donanimi = extract_ortam_donanimi(file_path)
        olcme_degerlendirme = extract_olcme_degerlendirme(file_path)
        kazanim_tablosu = extract_kazanim_tablosu(file_path)
        ogrenme_birimleri = extract_ogrenme_birimleri(file_path)
        uygulama_faaliyetleri = extract_uygulama_faaliyetleri(file_path)
        
        # İstatistikleri hesapla
        def kelime_sayisi(text):
            return len(text.split()) if text else 0
            
        def toplam_kazanim_saati(kazanim_tablosu):
            toplam_kazanim = sum(int(item.get('kazanim_sayisi', 0)) for item in kazanim_tablosu if item.get('kazanim_sayisi', '').isdigit())
            toplam_saat = sum(int(item.get('ders_saati_orani', 0)) for item in kazanim_tablosu if item.get('ders_saati_orani', '').isdigit())
            return toplam_kazanim, toplam_saat
        
        # İstatistikleri yazdır
        print(f"İşlenen Dosya          : {filename}")
        print(f"{'='*60}")
        print(f"Dersin Adı             : {ders_adi or 'Bulunamadı'}")
        print(f"Sınıf                  : {ders_sinifi or 'Bulunamadı'}")
        print(f"Dersin Süresi          : {ders_saati or 'Bulunamadı'}")
        print(f"Dersin Amacı           : {kelime_sayisi(dersin_amaci)} Kelime")
        print(f"Dersin Kazanımları     : {len(genel_kazanimlar)} Madde")
        print(f"EÖ Ortam ve Donanımı   : {len(ortam_donanimi)} Madde")
        print(f"Ölçme Değerlendirme    : {len(olcme_degerlendirme)} Madde")
        
        toplam_kazanim, toplam_saat = toplam_kazanim_saati(kazanim_tablosu)
        print(f"Dersin Kazanım Tablosu : {len(kazanim_tablosu)} Öğrenme Birimi, {toplam_kazanim} Kazanım, {toplam_saat} Saat")
        
        # ÖB - Konu - Kazanım detay istatistikleri
        toplam_ob = len(ogrenme_birimleri)
        toplam_konu = sum(len(birim.get('konular', [])) for birim in ogrenme_birimleri)
        toplam_detay_kazanim = sum(len(birim.get('kazanimlar', [])) for birim in ogrenme_birimleri)
        print(f"ÖB - Konu- Kazanım     : {toplam_ob} Öğrenme Birimi, {toplam_konu} Konu, {toplam_detay_kazanim} Kazanım")
        print(f"{'='*60}")
        print()
        
        # JSON yapısını oluştur
        result = {
            "metadata": {
                "processed_at": datetime.now().isoformat(),
                "source_file": filename,
                "file_path": file_path,
                "status": "success" if any([ders_adi, ders_sinifi, ders_saati, dersin_amaci]) else "partial"
            },
            "ders_bilgileri": {
                "ders_adi": ders_adi,
                "ders_sinifi": ders_sinifi,
                "haftalik_ders_saati": ders_saati,
                "dersin_amaci": dersin_amaci,
                "dersin_genel_kazanimlari": genel_kazanimlar,
                "egitim_ortam_donanimi": ortam_donanimi,
                "olcme_degerlendirme": olcme_degerlendirme
            },
            "ogrenme_birimleri_ozeti": kazanim_tablosu,
            "ogrenme_birimleri_detayi": ogrenme_birimleri,
            "uygulama_faaliyetleri": uygulama_faaliyetleri
        }
        
        return result
        
    except Exception as e:
        print(f"Error processing file: {str(e)}")
        return {
            "metadata": {
                "processed_at": datetime.now().isoformat(),
                "source_file": os.path.basename(file_path),
                "file_path": file_path,
                "status": "error",
                "error": str(e)
            },
            "ders_bilgileri": {},
            "ogrenme_birimleri_ozeti": [],
            "ogrenme_birimleri_detayi": [],
            "uygulama_faaliyetleri": []
        }

def process_directory(directory_path, output_json="oku_sonuc.json"):
    """
    Dizindeki tüm PDF ve DOCX dosyalarını işle ve sonuçları JSON'a kaydet
    """
    results = {}

    pdf_pattern = os.path.join(directory_path, "*.pdf")
    docx_pattern = os.path.join(directory_path, "*.docx")
    pdf_files = glob.glob(pdf_pattern)
    docx_files = glob.glob(docx_pattern)
    all_files = pdf_files + docx_files

    print(f"Found {len(pdf_files)} PDF and {len(docx_files)} DOCX files in {directory_path}")

    for file_path in all_files:
        filename = os.path.basename(file_path)
        result = oku(file_path)
        results[filename] = result

    # JSON'a kaydet
    output_data = {
        "metadata": {
            "processed_at": datetime.now().isoformat(),
            "source_directory": os.path.abspath(directory_path),
            "total_files": len(results),
            "successful_files": sum(1 for r in results.values() if r["metadata"]["status"] == "success"),
            "partial_files": sum(1 for r in results.values() if r["metadata"]["status"] == "partial"),
            "failed_files": sum(1 for r in results.values() if r["metadata"]["status"] == "error")
        },
        "results": results
    }

    try:
        with open(output_json, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)
        print(f"\n✓ Results saved to: {output_json}")
    except Exception as e:
        print(f"✗ Error saving JSON file: {str(e)}")

    return results

def main():
    """
    Ana fonksiyon - dizindeki PDF'leri işle
    """
    directory_path = "."  # Current directory
    output_json = "oku_sonuc.json"
    
    results = process_directory(directory_path, output_json)
    

if __name__ == "__main__":
    main()
